from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Pattern, Sequence


@dataclass(frozen=True)
class ParagraphRecord:
    index: int
    text: str


@dataclass(frozen=True)
class TableRecord:
    index: int
    rows: List[List[str]]


@dataclass(frozen=True)
class DocxContent:
    source_path: Path
    paragraphs: List[ParagraphRecord]
    tables: List[TableRecord]

    def paragraph(self, index: int) -> str:
        if index < 0 or index >= len(self.paragraphs):
            return ""
        return self.paragraphs[index].text

    def paragraphs_slice(self, start: int, end: int) -> List[ParagraphRecord]:
        low = max(0, start)
        high = min(len(self.paragraphs), end)
        return self.paragraphs[low:high]

    def table(self, index: int) -> TableRecord:
        for table in self.tables:
            if table.index == index:
                return table
        raise KeyError(f"table index {index} not found")


def _norm_text(text: str) -> str:
    return " ".join((text or "").split())


def load_docx(path: Path | str) -> DocxContent:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - runtime dependency guard
        raise RuntimeError("python-docx is required for CBA extraction") from exc

    source = Path(path)
    doc = Document(source)
    paragraphs = [ParagraphRecord(index=i, text=_norm_text(p.text)) for i, p in enumerate(doc.paragraphs)]
    tables: List[TableRecord] = []
    for i, table in enumerate(doc.tables):
        rows: List[List[str]] = []
        for row in table.rows:
            rows.append([_norm_text(c.text) for c in row.cells])
        tables.append(TableRecord(index=i, rows=rows))
    return DocxContent(source_path=source, paragraphs=paragraphs, tables=tables)


def find_paragraph(content: DocxContent, pattern: str | Pattern[str], *, start: int = 0) -> Optional[ParagraphRecord]:
    regex = re.compile(pattern, flags=re.IGNORECASE) if isinstance(pattern, str) else pattern
    for record in content.paragraphs[start:]:
        if regex.search(record.text):
            return record
    return None


def find_all_paragraphs(content: DocxContent, pattern: str | Pattern[str], *, start: int = 0) -> List[ParagraphRecord]:
    regex = re.compile(pattern, flags=re.IGNORECASE) if isinstance(pattern, str) else pattern
    return [record for record in content.paragraphs[start:] if regex.search(record.text)]


def parse_percent(text: str) -> float:
    m = re.search(r"(-?\d+(?:\.\d+)?)\s*%", text)
    if not m:
        raise ValueError(f"no percent in '{text}'")
    return float(m.group(1))


def parse_int(text: str) -> int:
    m = re.search(r"(-?\d+)", text)
    if not m:
        raise ValueError(f"no int in '{text}'")
    return int(m.group(1))


def parse_currency_amount(text: str) -> float:
    # Supports "$170 million", "$5 million", "$825,000", "$1.50-for-$1"
    cleaned = text.replace(",", "")
    m = re.search(r"\$(-?\d+(?:\.\d+)?)", cleaned)
    if not m:
        raise ValueError(f"no currency in '{text}'")
    value = float(m.group(1))
    if "million" in cleaned.lower():
        value *= 1_000_000.0
    return value


def extract_date_token(text: str) -> Optional[str]:
    # Normalized human-readable date snippets we use in rules.
    patterns = [
        r"(January\s+\d{1,2})",
        r"(March\s+\d{1,2})",
        r"(July\s+\d{1,2})",
        r"(December\s+\d{1,2})",
        r"(October\s+\d{1,2})",
        r"(September\s+\d{1,2})",
        r"(August\s+\d{1,2})",
    ]
    for pat in patterns:
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            return m.group(1)
    return None


def table_data_rows(table: TableRecord, header_rows: int = 1) -> Sequence[List[str]]:
    return table.rows[header_rows:]

